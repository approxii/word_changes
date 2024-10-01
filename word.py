import json
from docx import Document
from docx.oxml.ns import qn
import os
from docx.oxml import OxmlElement


class BookmarkExtractor:
    def __init__(self, docx_file):
        self.docx_file = docx_file #инициализиция файла
        self.templates = []  #массив всех шаблонов
        self.current_template = [] #текущий шаблон
        self.current_block = {}  #массив закладок
        self.template_count = 0  #кол-во шаблонов
        self.page_break_count = 0  #счетчик для разрыва страниц
        self.new_page_flag = False  #флаг для новой страницы

    def extract_and_save_bookmarks(self, json_file):
        doc = Document(self.docx_file) #инициализируем документ

        #проход по всем элементам(включая таблицы и тд)
        for element in doc.element.body.iter():
            if element.tag == qn('w:bookmarkStart'): #тег закладок для поиска в xml
                bookmark_name = element.get(qn('w:name'))
                if bookmark_name:
                    text = self.get_text_from_bookmarks(doc, element)
                    self.current_block[bookmark_name] = text

            #проверка на новую страницу
            if element.tag == qn('w:br') and element.get(qn('w:type')) == 'page':
                self.page_break_count += 1
                self.new_page_flag = True

                #логика для 1 разрыва страницы
                if self.page_break_count == 1 and self.current_block:
                    self.current_template.append(self.current_block)
                    self.current_block = {}

                #логика для 2 разрывов страницы
                if self.page_break_count == 2:
                    if self.current_block:
                        self.current_template.append(self.current_block)
                    self.template_count += 1
                    self.templates.append({
                        f"block {self.template_count}": self.current_template,
                        "newpage": str(self.new_page_flag).lower()
                    })
                    self.current_template = []
                    self.current_block = {}
                    self.page_break_count = 0
                    self.new_page_flag = False

        #для незакрытых блоков
        if self.current_block:
            self.current_template.append(self.current_block)
        if self.current_template:
            self.template_count += 1
            self.templates.append({
                f"block {self.template_count}": self.current_template,
                "newpage": "false"
            })

        output_data = {
            "blocks": self.templates  # Список всех шаблонов с массивами
        }
        self.save_to_json(output_data, json_file)

    #получаем текст с закладок
    def get_text_from_bookmarks(self, doc, bookmark_start):
        bookmark_text = []
        inside_bookmark = False
        start_id = bookmark_start.get(qn('w:id'))

        for element in doc.element.body.iter():
            #начало закладки
            if element.tag == qn('w:bookmarkStart') and element.get(qn('w:id')) == start_id:
                inside_bookmark = True

            #собираем текст между началом и концом закладки
            if inside_bookmark and element.tag == qn('w:t'):
                bookmark_text.append(element.text)

            #конец закладки
            if element.tag == qn('w:bookmarkEnd') and element.get(qn('w:id')) == start_id:
                #брейкаем после конца закладки
                break

        return ' '.join(bookmark_text).strip()

    #сохраняем в json
    def save_to_json(self, data, json_file):
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)


class DocxFiller:
    def __init__(self, docx_file, json_path, original_path):
        self.docx_file = docx_file
        self.json_path = json_path
        self.current_block = {}
        self.original_path = original_path
        self.original_folder = os.path.dirname(original_path)

    def load_json(self):
        #загрузка json'а
        with open(self.json_path, 'r', encoding='utf-8') as file:
            self.data = json.load(file)

    def update(self):
        #функция прохода по элементам документа и замены/добавления текста в местах закладок
        self.load_json()

        result_doc = Document()

        #проход по всем элементам(включая таблицы и тд)
        for index, block in enumerate(self.data['blocks']):
            doc, temp_filename = self.copy_to_temp(index)
            for key, value in block.items():
                print(f"Обработка блока {index}, ключ: {key}, значение: {value}")
                bookmark_found = False
                for element in doc.element.body.iter():
                    if element.tag == qn('w:bookmarkStart'):  # тег закладок для поиска в списке xml
                        bookmark_name = element.get(qn('w:name'))
                        if bookmark_name == key:
                            print(f"Закладка найдена: {bookmark_name}, текст для замены: {value}")
                            self.replace_text(doc, element, value)
                            bookmark_found = True
                if not bookmark_found:
                    print(f"Закладки в документе не найдены")

            doc.save(temp_filename)
            self.add_temp_to_original(result_doc, temp_filename)

            #удаление временных файлов
            if os.path.exists(temp_filename):
                os.remove(temp_filename)

        result_doc.save('result.docx')
        return result_doc

    def copy_to_temp(self, index):
        #функция копирования данных во временные файлы
        temp_filename = f'temp{index}.docx'
        original_doc = Document(self.docx_file)
        original_doc.save(temp_filename)
        return original_doc, temp_filename

    def add_temp_to_original(self, original_doc, temp_doc_path):
        #функция комбинирования временного файла с результатом
        temp_doc = Document(temp_doc_path)

        elements_to_copy = list(temp_doc.element.body)
        paragraph_index = 0
        table_index = 0

        for element in elements_to_copy:
            if element.tag.endswith('p'):
                if paragraph_index < len(temp_doc.paragraphs):
                    paragraph = temp_doc.paragraphs[paragraph_index]
                    self.copy_paragraph(paragraph, original_doc)
                    paragraph_index += 1
            elif element.tag.endswith('tbl'):
                if table_index < len(temp_doc.tables):
                    table = temp_doc.tables[table_index]

                    new_table = original_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style

                    for row_index, row in enumerate(table.rows):
                        for col_index, cell in enumerate(row.cells):
                            new_table.cell(row_index, col_index).text = cell.text

                            tcPr = new_table.cell(row_index, col_index)._tc.get_or_add_tcPr()
                            tcBorders = OxmlElement("w:tcBorders")
                            for border in ["top", "left", "bottom", "right"]:
                                element = OxmlElement(f"w:{border}")
                                element.set(qn("w:val"), "single")
                                element.set(qn("w:sz"), "4")
                                element.set(qn("w:space"), "0")
                                element.set(qn("w:color"), "auto")
                                tcBorders.append(element)
                            tcPr.append(tcBorders)

                    table_index += 1

        original_doc.add_page_break()


    def replace_text(self, doc, bookmark_element, new_text):
        #функция замены текста закладки на новый текст из json
        for sibling in bookmark_element.itersiblings():
            if sibling.tag == qn('w:r'):
                for child in sibling.iter():
                    if child.tag == qn('w:t'):
                        print(f"Изменили: {child.text} на {new_text}")
                        child.text = new_text
                        return


    def copy_paragraph(self, paragraph, document):
        #ункция для копирования абзаца в новый документ
        new_paragraph = document.add_paragraph()
        new_paragraph.style = paragraph.style  #стиль абзаца
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)  #текст
            new_run.bold = run.bold  #жирность
            new_run.italic = run.italic  #курсив
            new_run.font.size = run.font.size  #размер шрифта
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb  #цвет шрифта(если есть)




#main функция для тестов
def main():
    docx_file = 'ext.docx'  #изменить путь используя os
    json_file = 'bookmarks.json'  #изменить путь используя os

    extractor = BookmarkExtractor(docx_file)
    extractor.extract_and_save_bookmarks(json_file)

    doc = 'example.docx'
    json_path = 'data.json'

    doc_filler = DocxFiller(doc, json_path, doc)
    doc_filler.update()

main()
