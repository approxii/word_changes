from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from core.services.base import BaseDocumentService


class WordService(BaseDocumentService):
    def __init__(self):
        self.docx_file = None

    def load(self, file) -> None:
        self.docx_file = Document(file)


    def update(self, params: dict) -> None:
        if not self.docx_file:
            raise ValueError("Word файл не загружен.")
        
        result_doc = Document()

        #проход по всем элементам(включая таблицы и тд)
        for index, block in enumerate(params['blocks']):
            doc, temp_filename = self.copy_to_temp(index)
            for key, value in block.items():
                bookmark_found = False
                if isinstance(value, str):
                    for element in doc.element.body.iter():
                        #print(element.tag)
                        if element.tag == qn('w:bookmarkStart'):  # тег закладок для поиска в списке xml
                            bookmark_name = element.get(qn('w:name'))
                            if bookmark_name == key:
                                self.replace_text(doc, element, value)
                                bookmark_found = True
                    if not bookmark_found:
                        print(f"Закладка {key} в документе не найдена")
                elif isinstance(value, list):
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                bookmark_element = self.find_bookmark_in_table(cell, key)
                                if bookmark_element != None:
                                    cell.text = value[0]
                                    bookmark_found = True
                                    start_index = None
                                    for idx, c in enumerate(row.cells):
                                        if c.text == cell.text:
                                            start_index = idx
                                            break

                                    if start_index is not None:
                                        for i, val in enumerate(value[1:], start=1):
                                            if start_index + i < len(row.cells):
                                                row.cells[start_index + i].text = val

                                    break

                            if bookmark_found:
                                break
                        if bookmark_found:
                            break

                    if not bookmark_found:
                        print(f"Закладка '{key}' не найдена в документе.")
                    else:
                        print(f"Закладка '{key}' успешно обработана.")

            doc.save(temp_filename)
            self.add_temp_to_original(result_doc, temp_filename, params, index)

            #удаление временных файлов
            if os.path.exists(temp_filename):
                os.remove(temp_filename)

        if params['newpage'] == 'true':
            last_paragraph = result_doc.paragraphs[-1]
            p = last_paragraph._element
            p.getparent().remove(p)

        self.docx_file = result_doc

    def find_bookmark_in_table(self, cell, bookmark_name):
        for paragraph in cell.paragraphs:
            for element in paragraph._element.iter():
                if element.tag == qn('w:bookmarkStart'):
                    if element.get(qn('w:name')) == bookmark_name:
                        return element  #возвращаем элемент закладки
        return None  #если закладка не найдена

    def copy_to_temp(self, index):
        #функция копирования данных во временные файлы
        temp_filename = f'temp{index}.docx'

        self.docx_file.save(temp_filename)

        new_doc = Document(temp_filename)

        return new_doc, temp_filename

    def add_temp_to_original(self, original_doc, temp_doc_path, params: dict, index):
        #функция комбинирования временного файла с результатом
        temp_doc = Document(temp_doc_path)

        elements_to_copy = list(temp_doc.element.body)
        paragraph_index = 0
        table_index = 0

        for element in elements_to_copy:
            if element.tag.endswith('p'):
                if paragraph_index < len(temp_doc.paragraphs):
                    paragraph = temp_doc.paragraphs[paragraph_index]
                    self.copy_paragraph(paragraph, original_doc)
                    paragraph_index += 1
            elif element.tag.endswith('tbl'):
                if table_index < len(temp_doc.tables):
                    table = temp_doc.tables[table_index]

                    new_table = original_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style

                    for row_index, row in enumerate(table.rows):
                        for col_index, cell in enumerate(row.cells):
                            new_table.cell(row_index, col_index).text = cell.text

                            tcPr = new_table.cell(row_index, col_index)._tc.get_or_add_tcPr()
                            tcBorders = OxmlElement("w:tcBorders")
                            for border in ["top", "left", "bottom", "right"]:
                                element = OxmlElement(f"w:{border}")
                                element.set(qn("w:val"), "single")
                                element.set(qn("w:sz"), "4")
                                element.set(qn("w:space"), "0")
                                element.set(qn("w:color"), "auto")
                                tcBorders.append(element)
                            tcPr.append(tcBorders)

                    table_index += 1
        if params['newpage'] == 'true':
            original_doc.add_page_break()

    def replace_text(self, doc, bookmark_element, new_text):
        parent_element = bookmark_element.getparent()

        for sibling in bookmark_element.itersiblings():
            if sibling.tag == qn('w:r'):
                text_elements = sibling.findall(qn('w:t'))

                if text_elements:
                    for child in text_elements:
                        child.text = new_text  # Замена на новый текст
                        return

        new_run = OxmlElement('w:r')
        new_text_element = OxmlElement('w:t')
        new_text_element.text = new_text

        new_run.append(new_text_element)
        parent_element.insert(parent_element.index(bookmark_element), new_run)
        parent_element.remove(bookmark_element)



    def copy_paragraph(self, paragraph, document):
        #ункция для копирования абзаца в новый документ
        new_paragraph = document.add_paragraph()
        new_paragraph.style = paragraph.style  #стиль абзаца
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)  #текст
            new_run.bold = run.bold  #жирность
            new_run.italic = run.italic  #курсив
            new_run.font.size = run.font.size  #размер шрифта
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb  #цвет шрифта(если есть)

    def save_to_bytes(self) -> BytesIO:
        if not self.docx_file:
            raise ValueError("Word файл не загружен.")
        output = BytesIO()
        self.docx_file.save(output)
        output.seek(0)
        return output

    def save_to_file(self, file_path: str) -> None:
        if self.docx_file:
            self.docx_file.save(file_path)
        else:
            raise ValueError("Word файл не загружен.")